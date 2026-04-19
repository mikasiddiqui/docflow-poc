from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph


def generate_docx(template_path: Path, structured_content: dict, output_path: Path) -> None:
    document = Document(template_path)
    anchor = _find_placeholder_paragraph(document, "{{items_table}}")
    replacements = {
        "{{source_file_name}}": structured_content.get("source_file_name", ""),
        "{{generated_at}}": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "{{items_count}}": str(len(structured_content.get("extracted_items", []))),
    }

    _replace_placeholders(document, replacements)
    _insert_items_table(document, structured_content.get("extracted_items", []), anchor)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _replace_placeholders(document: DocumentObject, replacements: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)


def _replace_in_paragraph(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    for run in paragraph.runs:
        for placeholder, value in replacements.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)


def _insert_items_table(
    document: DocumentObject,
    extracted_items: list[dict],
    anchor: Paragraph | None,
) -> None:
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = ["Page", "Type", "Content", "Author", "Subject"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    if extracted_items:
        for item in extracted_items:
            row = table.add_row().cells
            row[0].text = str(item.get("page", ""))
            row[1].text = str(item.get("type", ""))
            row[2].text = str(item.get("content", ""))
            row[3].text = str(item.get("author", ""))
            row[4].text = str(item.get("subject", ""))
    else:
        row = table.add_row().cells
        row[0].text = ""
        row[1].text = "note"
        row[2].text = "No extracted items were available."
        row[3].text = ""
        row[4].text = ""

    if anchor is not None:
        anchor._p.addnext(table._tbl)
        _remove_paragraph(anchor)


def _find_placeholder_paragraph(document: DocumentObject, placeholder: str) -> Paragraph | None:
    for paragraph in document.paragraphs:
        if placeholder in paragraph.text:
            return paragraph
    return None


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
